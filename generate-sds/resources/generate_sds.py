"""
WHOOP SDS Document Generator

Opens whoop_sds_template.docx, clears its body, and fills it from a
content JSON file — preserving all styles, fonts, heading colors, and
page layout from the template.

Usage:
    python3 generate_sds.py content.json output.docx [template.docx]

content.json schema: see SKILL.md → CONTENT SCHEMA
"""

import json
import sys
import io
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches


W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# ── Formatting constants (WHOOP SaMD standard) ────────────────────────────────
TBL_HEADER       = 'F2F2F2'   # table header row fill
TBL_ALT          = 'EFEFEF'   # alternating data row fill
CODE_BG          = 'F6F8FA'   # code/JSON block background
TBL_BORDER_COLOR = '000000'   # table border color (black)
H1_COLOR         = RGBColor(0x2F, 0x54, 0x96)   # heading 1/2 blue
H3_COLOR         = RGBColor(0x1F, 0x37, 0x63)   # heading 3 navy
CAPTION_COLOR    = RGBColor(0x2F, 0x54, 0x96)


# ── XML helpers ────────────────────────────────────────────────────────────────
def _wns(tag):
    return f'{{{W}}}{tag}'


def _set_cell_shd(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(_wns('shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)


def _add_tbl_borders(table, color=TBL_BORDER_COLOR):
    tbl_el = table._tbl
    tblPr = tbl_el.find(_wns('tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl_el.insert(0, tblPr)
    for old in tblPr.findall(_wns('tblBorders')):
        tblPr.remove(old)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _cell_run(cell_para, text, bold=False, italic=False, size_pt=9,
              font='Arial', color_rgb=None):
    r = cell_para.add_run(text)
    r.font.name = font
    r.font.size = Pt(size_pt)
    r.font.bold = bold
    r.font.italic = italic
    if color_rgb:
        r.font.color.rgb = color_rgb
    return r


# ── Document generator ─────────────────────────────────────────────────────────
class SDSGenerator:
    def __init__(self, template_path: str):
        self.doc = Document(template_path)
        self._clear_body()
        self._accept_tracked_changes()

    # ── Template prep ─────────────────────────────────────────────────────────
    def _clear_body(self):
        """Remove all body content; keep sectPr (page settings, margins, headers)."""
        body = self.doc.element.body
        for child in list(body):
            if child.tag != _wns('sectPr'):
                body.remove(child)

    def _accept_tracked_changes(self):
        """Accept all tracked changes across every XML part (body, headers, footers).

        w:del  → remove (accept deletion)
        w:ins  → unwrap, keep children (accept insertion)
        w:*PrChange → remove (keep current formatting)
        """
        for part in self.doc.part.package.iter_parts():
            try:
                el = part._element
            except AttributeError:
                continue
            if el is None:
                continue
            for del_el in el.findall(f'.//{_wns("del")}'):
                parent = del_el.getparent()
                if parent is not None:
                    parent.remove(del_el)
            for ins_el in el.findall(f'.//{_wns("ins")}'):
                parent = ins_el.getparent()
                if parent is None:
                    continue
                idx = list(parent).index(ins_el)
                for child in list(ins_el):
                    parent.insert(idx, child)
                    idx += 1
                parent.remove(ins_el)
            for tag in (_wns('rPrChange'), _wns('pPrChange'), _wns('tblPrChange'),
                        _wns('trPrChange'), _wns('tcPrChange'), _wns('sectPrChange')):
                for el2 in el.findall(f'.//{tag}'):
                    parent = el2.getparent()
                    if parent is not None:
                        parent.remove(el2)

    # ── Page header (WHOOP standard 3-row table in header1.xml) ──────────────
    def _update_header(self, doc_id: str, title: str, version: str):
        """Fill Row 2 of the WHOOP header table: doc_id | title | version."""
        try:
            header = self.doc.sections[0].header
            tbl = header.tables[0]
            values_row = tbl.rows[2]
            for cell, text in zip(values_row.cells, [doc_id, title, version]):
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.text = ''
                para = cell.paragraphs[0]
                if para.runs:
                    para.runs[0].text = text
                else:
                    r = para.add_run(text)
                    r.font.name = 'Arial'
                    r.font.size = Pt(10)
        except Exception as e:
            print(f'Warning: could not update page header: {e}')

    # ── Title block ───────────────────────────────────────────────────────────
    def add_title_block(self, doc_id: str, title: str, version: str, date: str):
        self._update_header(doc_id, title, version)

        p = self.doc.add_paragraph()
        r = p.add_run(doc_id)
        r.font.name = 'Arial'; r.font.size = Pt(10)
        r.font.color.rgb = H1_COLOR
        p.paragraph_format.space_after = Pt(2)

        p2 = self.doc.add_paragraph()
        r2 = p2.add_run(title)
        r2.font.name = 'Arial'; r2.font.size = Pt(20); r2.font.bold = True
        r2.font.color.rgb = H3_COLOR
        p2.paragraph_format.space_after = Pt(4)

        p3 = self.doc.add_paragraph()
        r3 = p3.add_run(f'Version {version}  ·  {date}')
        r3.font.name = 'Arial'; r3.font.size = Pt(10); r3.font.italic = True
        r3.font.color.rgb = H1_COLOR
        p3.paragraph_format.space_after = Pt(12)

        # Blue horizontal rule
        p_hr = self.doc.add_paragraph()
        pPr = p_hr._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), '2F5496')
        pBdr.append(bottom); pPr.append(pBdr)
        p_hr.paragraph_format.space_after = Pt(8)

    # ── Approvers table ───────────────────────────────────────────────────────
    def add_approvers(self, approvers: list):
        p = self.doc.add_paragraph()
        r = p.add_run('Approvers Table:')
        r.font.name = 'Arial'; r.font.size = Pt(10); r.font.bold = True
        p.paragraph_format.space_after = Pt(2)

        tbl = self.doc.add_table(rows=1 + len(approvers), cols=2)
        tbl.style = 'Normal Table'
        _add_tbl_borders(tbl)
        tbl.columns[0].width = Emu(2155 * 914)
        tbl.columns[1].width = Emu(7195 * 914)

        for cell, txt in zip(tbl.rows[0].cells, ['Role/Function', 'Name']):
            _set_cell_shd(cell, TBL_HEADER)
            cell.paragraphs[0].clear()
            _cell_run(cell.paragraphs[0], txt, bold=True)

        for i, appr in enumerate(approvers, 1):
            row = tbl.rows[i].cells
            row[0].paragraphs[0].clear()
            row[1].paragraphs[0].clear()
            _cell_run(row[0].paragraphs[0], appr.get('role', ''))
            _cell_run(row[1].paragraphs[0], appr.get('name', ''))
            if i % 2 == 0:
                _set_cell_shd(row[0], TBL_ALT)
                _set_cell_shd(row[1], TBL_ALT)
        self.doc.add_paragraph()

    # ── Table of Contents ─────────────────────────────────────────────────────
    def add_toc(self, sections: list):
        try:
            p_title = self.doc.add_paragraph(style='TOC Heading')
            p_title.clear()
            r = p_title.add_run('Table of Contents')
            r.font.name = 'Arial'; r.font.size = Pt(12); r.font.bold = True
        except Exception:
            p_title = self.doc.add_paragraph('Table of Contents')
            for run in p_title.runs:
                run.font.name = 'Arial'; run.font.size = Pt(12); run.font.bold = True
        p_title.paragraph_format.space_after = Pt(6)

        for sec in sections:
            p = self.doc.add_paragraph(style='toc 1')
            p.clear()
            r = p.add_run(f"{sec['number']}  {sec['title']}")
            r.font.name = 'Arial'; r.font.size = Pt(10)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            for sub in sec.get('subsections', []):
                ps = self.doc.add_paragraph(style='toc 2')
                ps.clear()
                rs = ps.add_run(f"    {sec['number']}.{sub['number']}  {sub['title']}")
                rs.font.name = 'Arial'; rs.font.size = Pt(9)
                ps.paragraph_format.space_before = Pt(0)
                ps.paragraph_format.space_after = Pt(1)
        self.doc.add_page_break()

    # ── Section heading ───────────────────────────────────────────────────────
    def add_section_heading(self, number: str, title: str, level: int = 1):
        label = f'{number}  {title}' if number else title
        h = self.doc.add_heading(label, level=level)
        for run in h.runs:
            run.font.name = 'Arial'

    # ── Body paragraph ────────────────────────────────────────────────────────
    def add_paragraph(self, text: str):
        p = self.doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = 'Arial'; r.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(6)

    # ── Data table ────────────────────────────────────────────────────────────
    def add_table(self, headers: list, rows: list,
                  caption: str = None, col_widths: list = None):
        if caption:
            cp = self.doc.add_paragraph()
            r = cp.add_run(caption)
            r.font.name = 'Arial'; r.font.size = Pt(10)
            r.font.bold = True; r.font.italic = True
            cp.paragraph_format.space_after = Pt(2)

        tbl = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.style = 'Normal Table'
        _add_tbl_borders(tbl)

        # Column widths: total text area = 9350 dxa (6.5 inch)
        n = len(headers)
        if col_widths and len(col_widths) == n:
            dxa_widths = [int(w * 1440) for w in col_widths]
        else:
            dxa_widths = [9350 // n] * n
        for i, col in enumerate(tbl.columns):
            col.width = Emu(dxa_widths[i] * 914)

        for cell, txt in zip(tbl.rows[0].cells, headers):
            _set_cell_shd(cell, TBL_HEADER)
            cell.paragraphs[0].clear()
            _cell_run(cell.paragraphs[0], txt, bold=True)

        for i, row_data in enumerate(rows):
            row = tbl.rows[i + 1].cells
            for cell, val in zip(row, row_data):
                cell.paragraphs[0].clear()
                _cell_run(cell.paragraphs[0], str(val))
                if i % 2 == 0:
                    _set_cell_shd(cell, TBL_ALT)
        self.doc.add_paragraph()

    # ── Code / JSON block ─────────────────────────────────────────────────────
    def add_code_block(self, code: str, label: str = None):
        if label:
            lp = self.doc.add_paragraph()
            r = lp.add_run(label)
            r.font.name = 'Arial'; r.font.size = Pt(9); r.font.italic = True
            r.font.color.rgb = CAPTION_COLOR
            lp.paragraph_format.space_after = Pt(1)

        tbl = self.doc.add_table(rows=1, cols=1)
        tbl.style = 'Normal Table'
        _add_tbl_borders(tbl)
        tbl.columns[0].width = Emu(9350 * 914)

        cell = tbl.rows[0].cells[0]
        _set_cell_shd(cell, CODE_BG)
        first = True
        for line in code.strip().split('\n'):
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            r = p.add_run(line if line else ' ')
            r.font.name = 'Courier New'; r.font.size = Pt(8.5)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        self.doc.add_paragraph()

    # ── Figures ───────────────────────────────────────────────────────────────
    def embed_image_file(self, path: str, caption: str = None, width_in: float = 5.5):
        """Embed a local PNG/JPEG (from Confluence download or any source)."""
        img_path = Path(path).expanduser()
        if not img_path.exists() or img_path.stat().st_size < 100:
            print(f'Warning: image missing or empty: {img_path}')
            p = self.doc.add_paragraph(f'[IMAGE NOT AVAILABLE: {img_path.name}]')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if caption:
                self._add_figure_caption(caption)
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(img_path), width=Inches(width_in))
        if caption:
            self._add_figure_caption(caption)

    def embed_matplotlib_figure(self, fig, caption: str = None, width_in: float = 5.5):
        """Embed a matplotlib figure (for auto-generated diagrams)."""
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
        buf.seek(0)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(buf, width=Inches(width_in))
        plt.close(fig)
        if caption:
            self._add_figure_caption(caption)

    def _add_figure_caption(self, caption: str):
        cp = self.doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        r.font.name = 'Arial'; r.font.size = Pt(9); r.font.italic = True
        r.font.color.rgb = CAPTION_COLOR
        cp.paragraph_format.space_after = Pt(8)

    # ── Revision history ──────────────────────────────────────────────────────
    def add_revision_history(self, history: list):
        self.add_section_heading('', 'Document Revision History', level=1)
        headers = ['Version', 'Date', 'Description', 'Author']
        rows = [[h.get('version',''), h.get('date',''),
                 h.get('description',''), h.get('author','')]
                for h in history]
        self.add_table(headers, rows, col_widths=[0.8, 1.0, 3.7, 1.0])

    def add_page_break(self):
        self.doc.add_page_break()

    def save(self, path: str):
        self.doc.save(path)
        print(f'Saved: {path}')


# ── Auto-generated diagram helpers ────────────────────────────────────────────
def make_architecture_diagram(title: str, components: list, flows: list):
    """Generate a component-flow architecture diagram."""
    fig, ax = plt.subplots(figsize=(9, 4.5))
    ax.set_xlim(0, 10); ax.set_ylim(0, 5); ax.axis('off')
    ax.set_title(title, fontsize=12, fontweight='bold', color='#2F5496', pad=12)

    boxes = []
    for comp in components:
        x, y = comp['x'], comp['y']
        box = mpatches.FancyBboxPatch(
            (x - 1.2, y - 0.45), 2.4, 0.9,
            boxstyle='round,pad=0.08', linewidth=1.5,
            edgecolor='#1F3763', facecolor=comp.get('color', '#2E75B6'), zorder=3)
        ax.add_patch(box)
        ax.text(x, y + 0.1, comp['label'], ha='center', va='center',
                fontsize=8.5, fontweight='bold', color='white', zorder=4)
        if comp.get('sub'):
            ax.text(x, y - 0.2, comp['sub'], ha='center', va='center',
                    fontsize=6.5, color='#DDDDDD', zorder=4)
        boxes.append((x, y))

    for flow in flows:
        x1, y1 = boxes[flow['from']]; x2, y2 = boxes[flow['to']]
        dx, dy = x2 - x1, y2 - y1
        s = 0.82
        mx = x1 + dx * (1 - s) / 2; my = y1 + dy * (1 - s) / 2
        ex = x1 + dx * (1 - (1 - s) / 2); ey = y1 + dy * (1 - (1 - s) / 2)
        ax.annotate('', xy=(ex, ey), xytext=(mx, my),
                    arrowprops=dict(arrowstyle='->', color='#1F3763', lw=1.5))
        if flow.get('label'):
            ax.text((mx + ex) / 2, (my + ey) / 2 + 0.12, flow['label'],
                    ha='center', va='bottom', fontsize=7, color='#444444',
                    bbox=dict(boxstyle='round,pad=0.1', facecolor='white',
                              edgecolor='none', alpha=0.9))
    fig.tight_layout()
    return fig


def make_state_machine_diagram(title: str, states: list, transitions: list):
    """Generate a horizontal state machine diagram."""
    n = len(states)
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.set_xlim(-0.5, n + 0.5); ax.set_ylim(-1.5, 2); ax.axis('off')
    ax.set_title(title, fontsize=11, fontweight='bold', color='#2F5496', pad=10)

    pos = {s: (i * (9.5 / (n - 1)) if n > 1 else 4.5, 0)
           for i, s in enumerate(states)}
    DEFAULT_COLORS = ['#BDD7EE', '#9DC3E6', '#2E75B6', '#1F3763',
                      '#C00000', '#70AD47', '#ED7D31', '#9E9E9E']
    state_colors = {s: DEFAULT_COLORS[i % len(DEFAULT_COLORS)]
                    for i, s in enumerate(states)}

    for state, (x, y) in pos.items():
        color = state_colors[state]
        ax.add_patch(mpatches.FancyBboxPatch(
            (x - 1.1, y - 0.35), 2.2, 0.7,
            boxstyle='round,pad=0.07', linewidth=1.5,
            edgecolor='#1F3763', facecolor=color, zorder=3))
        tc = 'white' if color not in ('#BDD7EE', '#9DC3E6') else '#1F3763'
        ax.text(x, y, state.replace('_', '\n'), ha='center', va='center',
                fontsize=6.5, fontweight='bold', color=tc, zorder=4)

    for t in transitions:
        if t['from'] in pos and t['to'] in pos:
            x1, y1 = pos[t['from']]; x2, y2 = pos[t['to']]
            off = 0.25 if y1 == y2 else 0
            ax.annotate('', xy=(x2 - 1.1, y2 + off), xytext=(x1 + 1.1, y1 + off),
                        arrowprops=dict(arrowstyle='->', color='#555555', lw=1.2,
                                        connectionstyle=f'arc3,rad={0.3 if off else 0}'))
            ax.text((x1 + x2) / 2, (y1 + y2) / 2 + (0.5 if off else 0.18),
                    t.get('label', ''), ha='center', va='bottom',
                    fontsize=6, color='#333333',
                    bbox=dict(boxstyle='round,pad=0.05', facecolor='white',
                              edgecolor='none', alpha=0.85))
    fig.tight_layout()
    return fig


# ── Section renderer ──────────────────────────────────────────────────────────
def _render(gen: SDSGenerator, sec: dict):
    for para in sec.get('paragraphs', []):
        gen.add_paragraph(para)
    for fig_def in sec.get('figures', []):
        width = fig_def.get('width', 5.5)
        fig_type = fig_def.get('type', '')
        if fig_type == 'image':
            gen.embed_image_file(
                fig_def['path'], fig_def.get('caption'), width_in=width)
        elif fig_type == 'architecture':
            gen.embed_matplotlib_figure(
                make_architecture_diagram(
                    fig_def.get('title', ''),
                    fig_def['components'],
                    fig_def['flows']),
                fig_def.get('caption'), width_in=width)
        elif fig_type == 'state_machine':
            gen.embed_matplotlib_figure(
                make_state_machine_diagram(
                    fig_def.get('title', ''),
                    fig_def['states'],
                    fig_def['transitions']),
                fig_def.get('caption'), width_in=width)
    for tbl in sec.get('tables', []):
        gen.add_table(
            tbl['headers'], tbl['rows'],
            caption=tbl.get('caption'),
            col_widths=tbl.get('col_widths'))
    for cb in sec.get('code_blocks', []):
        gen.add_code_block(cb['code'], label=cb.get('label'))


# ── Entry point ───────────────────────────────────────────────────────────────
def main():
    if len(sys.argv) < 3:
        print('Usage: generate_sds.py content.json output.docx [template.docx]')
        sys.exit(1)

    content_path = sys.argv[1]
    output_path  = sys.argv[2]
    template_path = (sys.argv[3] if len(sys.argv) >= 4
                     else str(Path(__file__).parent / 'whoop_sds_template.docx'))

    if not Path(template_path).exists():
        print(f'ERROR: Template not found: {template_path}')
        sys.exit(1)

    with open(content_path) as f:
        content = json.load(f)

    gen = SDSGenerator(template_path)

    gen.add_title_block(
        content['doc_id'], content['title'],
        content['algo_version'], content['date'])

    if content.get('approvers'):
        gen.add_approvers(content['approvers'])

    gen.add_page_break()
    gen.add_toc(content.get('sections', []))

    for sec in content.get('sections', []):
        gen.add_section_heading(sec['number'], sec['title'], level=1)
        _render(gen, sec)
        for sub in sec.get('subsections', []):
            gen.add_section_heading(
                f"{sec['number']}.{sub['number']}", sub['title'], level=2)
            _render(gen, sub)
            for subsub in sub.get('subsections', []):
                gen.add_section_heading(
                    f"{sec['number']}.{sub['number']}.{subsub['number']}",
                    subsub['title'], level=3)
                _render(gen, subsub)

    if content.get('revision_history'):
        gen.add_page_break()
        gen.add_revision_history(content['revision_history'])

    gen.save(output_path)


if __name__ == '__main__':
    main()
