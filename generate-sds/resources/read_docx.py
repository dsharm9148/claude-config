"""
Extract structured plain text from a .docx file for validation comparison.

Preserves section heading markers so the Validation Agent can compare documents
section-by-section rather than as an undifferentiated blob.

Usage:
    python3 read_docx.py <path/to/document.docx> [> output.txt]

Output format:
    ## Heading 1 text
    ### Heading 2 text
    #### Heading 3 text
    Normal paragraph text...
    [TABLE]
    col1\tcol2\tcol3
    val1\tval2\tval3
    [/TABLE]
"""

import sys
import pathlib
from docx import Document
from docx.oxml.ns import qn


def _is_heading(para, level: int) -> bool:
    style_name = para.style.name if para.style else ''
    return style_name.lower().startswith(f'heading {level}')


def _is_any_heading(para) -> bool:
    style_name = para.style.name if para.style else ''
    return style_name.lower().startswith('heading')


def _get_heading_level(para) -> int:
    style_name = (para.style.name or '').lower()
    for level in range(1, 7):
        if style_name.startswith(f'heading {level}'):
            return level
    return 0


def _para_text(para) -> str:
    """Get full text of a paragraph including all run text."""
    return ''.join(run.text for run in para.runs) or para.text


def _table_text(table) -> list[str]:
    """Render a table as tab-separated rows."""
    lines = ['[TABLE]']
    for row in table.rows:
        cells = []
        for cell in row.cells:
            # Get plain text from cell paragraphs
            cell_text = ' '.join(
                p.text.strip() for p in cell.paragraphs if p.text.strip()
            )
            cells.append(cell_text)
        if any(cells):
            lines.append('\t'.join(cells))
    lines.append('[/TABLE]')
    return lines


def extract_document(doc_path: str) -> str:
    """
    Extract structured text from a docx file.

    Returns a string with heading markers, paragraph text, and table markers
    interleaved in document order.
    """
    doc = Document(doc_path)
    output_lines = []

    # Walk the document body in order to interleave paragraphs and tables
    body = doc.element.body
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            # Find the matching paragraph object by xml element
            para = None
            for p in doc.paragraphs:
                if p._element is child:
                    para = p
                    break
            if para is None:
                continue

            text = _para_text(para).strip()
            if not text:
                continue

            level = _get_heading_level(para)
            if level == 1:
                output_lines.append(f'\n## {text}')
            elif level == 2:
                output_lines.append(f'\n### {text}')
            elif level == 3:
                output_lines.append(f'\n#### {text}')
            elif level >= 4:
                output_lines.append(f'\n##### {text}')
            else:
                output_lines.append(text)

        elif tag == 'tbl':
            # Find the matching table object
            tbl = None
            for t in doc.tables:
                if t._element is child:
                    tbl = t
                    break
            if tbl is None:
                continue
            output_lines.extend(_table_text(tbl))

    # Collapse runs of more than 2 blank lines to 1
    cleaned = []
    blank_count = 0
    for line in output_lines:
        if line.strip() == '':
            blank_count += 1
            if blank_count <= 1:
                cleaned.append(line)
        else:
            blank_count = 0
            cleaned.append(line)

    return '\n'.join(cleaned)


def main():
    if len(sys.argv) < 2:
        print('Usage: python3 read_docx.py <path/to/document.docx>', file=sys.stderr)
        sys.exit(1)

    doc_path = pathlib.Path(sys.argv[1]).expanduser()
    if not doc_path.exists():
        print(f'Error: file not found: {doc_path}', file=sys.stderr)
        sys.exit(1)

    print(extract_document(str(doc_path)))


if __name__ == '__main__':
    main()
